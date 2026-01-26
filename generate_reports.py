#!/usr/bin/env python3
"""
Enhanced Report Generator
Generates DOCX and XLSX reports from scan JSON results
Requires: python-docx, openpyxl
Install: pip install python-docx openpyxl
"""

import json
import sys
import argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[!] Warning: python-docx not installed. DOCX generation disabled.")
    print("    Install with: pip install python-docx")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    print("[!] Warning: openpyxl not installed. XLSX generation disabled.")
    print("    Install with: pip install openpyxl")


class EnhancedReportGenerator:
    def __init__(self, json_file: str):
        self.json_file = Path(json_file)
        self.data = self.load_json()
        self.output_dir = self.json_file.parent
        
    def load_json(self):
        """Load JSON scan results"""
        print(f"[*] Loading results from {self.json_file}")
        with open(self.json_file, 'r') as f:
            return json.load(f)
    
    def generate_docx_report(self):
        """Generate a professional DOCX report"""
        if not DOCX_AVAILABLE:
            print("[!] Cannot generate DOCX report - python-docx not installed")
            return None
        
        print("[*] Generating DOCX report...")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Penetration Test Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading('Executive Summary', 1)
        
        summary_table = doc.add_table(rows=4, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        summary_data = [
            ('Scan Date', self.data.get('scan_date', 'N/A')),
            ('Target Networks', ', '.join(self.data.get('targets', []))),
            ('Live Hosts Discovered', str(len(self.data.get('live_hosts', [])))),
            ('Hosts with Open Ports', str(len(self.data.get('detailed_scans', {}))))
        ]
        
        for idx, (label, value) in enumerate(summary_data):
            row = summary_table.rows[idx]
            row.cells[0].text = label
            row.cells[1].text = value
            # Bold the labels
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Live Hosts Section
        doc.add_heading('Discovered Live Hosts', 1)
        
        if self.data.get('live_hosts'):
            for host in sorted(self.data['live_hosts']):
                doc.add_paragraph(f'• {host}', style='List Bullet')
        else:
            doc.add_paragraph('No live hosts discovered.')
        
        doc.add_page_break()
        
        # Detailed Findings
        doc.add_heading('Detailed Findings', 1)
        
        detailed_scans = self.data.get('detailed_scans', {})
        
        if not detailed_scans:
            doc.add_paragraph('No detailed scan data available.')
        else:
            for host in sorted(detailed_scans.keys()):
                host_data = detailed_scans[host]
                
                # Host heading
                doc.add_heading(f'Host: {host}', 2)
                
                services = host_data.get('services', [])
                
                if not services:
                    doc.add_paragraph('No services detected.')
                    continue
                
                # Services table
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'
                
                # Header row
                header_cells = table.rows[0].cells
                headers = ['Port', 'Protocol', 'Service', 'Version']
                for idx, header in enumerate(headers):
                    header_cells[idx].text = header
                    header_cells[idx].paragraphs[0].runs[0].font.bold = True
                
                # Service rows
                for service in services:
                    row_cells = table.add_row().cells
                    row_cells[0].text = service.get('port', 'N/A')
                    row_cells[1].text = service.get('protocol', 'N/A')
                    row_cells[2].text = service.get('service', 'N/A')
                    row_cells[3].text = service.get('version', 'N/A')
                
                # Script outputs
                for service in services:
                    if service.get('scripts'):
                        doc.add_paragraph()
                        doc.add_paragraph(
                            f"Port {service['port']} Script Results:",
                            style='Intense Quote'
                        )
                        
                        for script in service['scripts']:
                            doc.add_paragraph(f"• {script['id']}", style='List Bullet')
                            
                            # Add script output in monospace
                            output_para = doc.add_paragraph(script['output'])
                            output_para.style = 'No Spacing'
                            for run in output_para.runs:
                                run.font.name = 'Courier New'
                                run.font.size = Pt(9)
                
                doc.add_paragraph()  # Spacing between hosts
        
        # Save document
        output_file = self.output_dir / f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(output_file)
        
        print(f"[+] DOCX report saved: {output_file}")
        return output_file
    
    def generate_xlsx_report(self):
        """Generate an Excel spreadsheet report"""
        if not XLSX_AVAILABLE:
            print("[!] Cannot generate XLSX report - openpyxl not installed")
            return None
        
        print("[*] Generating XLSX report...")
        
        wb = Workbook()
        
        # Remove default sheet
        wb.remove(wb.active)
        
        # Summary Sheet
        ws_summary = wb.create_sheet("Summary")
        
        # Header styling
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True, size=12)
        
        # Add summary data
        ws_summary['A1'] = 'Penetration Test Report Summary'
        ws_summary['A1'].font = Font(bold=True, size=14)
        ws_summary.merge_cells('A1:B1')
        
        summary_data = [
            ('Scan Date', self.data.get('scan_date', 'N/A')),
            ('Target Networks', ', '.join(self.data.get('targets', []))),
            ('Live Hosts', len(self.data.get('live_hosts', []))),
            ('Hosts with Services', len(self.data.get('detailed_scans', {})))
        ]
        
        row = 3
        for label, value in summary_data:
            ws_summary[f'A{row}'] = label
            ws_summary[f'B{row}'] = value
            ws_summary[f'A{row}'].font = Font(bold=True)
            row += 1
        
        # Auto-size columns
        ws_summary.column_dimensions['A'].width = 20
        ws_summary.column_dimensions['B'].width = 40
        
        # Live Hosts Sheet
        ws_hosts = wb.create_sheet("Live Hosts")
        
        ws_hosts['A1'] = 'IP Address'
        ws_hosts['B1'] = 'Open Ports Count'
        ws_hosts['A1'].fill = header_fill
        ws_hosts['A1'].font = header_font
        ws_hosts['B1'].fill = header_fill
        ws_hosts['B1'].font = header_font
        
        row = 2
        for host in sorted(self.data.get('live_hosts', [])):
            ws_hosts[f'A{row}'] = host
            
            # Count open ports
            port_count = 0
            if host in self.data.get('detailed_scans', {}):
                port_count = len(self.data['detailed_scans'][host].get('services', []))
            
            ws_hosts[f'B{row}'] = port_count
            row += 1
        
        ws_hosts.column_dimensions['A'].width = 20
        ws_hosts.column_dimensions['B'].width = 20
        
        # Detailed Services Sheet
        ws_services = wb.create_sheet("Services")
        
        headers = ['Host', 'Port', 'Protocol', 'State', 'Service', 'Version']
        for col, header in enumerate(headers, 1):
            cell = ws_services.cell(1, col, header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center')
        
        row = 2
        for host in sorted(self.data.get('detailed_scans', {}).keys()):
            host_data = self.data['detailed_scans'][host]
            services = host_data.get('services', [])
            
            for service in services:
                ws_services.cell(row, 1, host)
                ws_services.cell(row, 2, service.get('port', 'N/A'))
                ws_services.cell(row, 3, service.get('protocol', 'N/A'))
                ws_services.cell(row, 4, service.get('state', 'N/A'))
                ws_services.cell(row, 5, service.get('service', 'N/A'))
                ws_services.cell(row, 6, service.get('version', 'N/A'))
                row += 1
        
        # Auto-size columns
        for col in range(1, 7):
            ws_services.column_dimensions[chr(64 + col)].width = 20
        
        # Vulnerabilities/Findings Sheet (from scripts)
        ws_findings = wb.create_sheet("Script Findings")
        
        headers = ['Host', 'Port', 'Script ID', 'Finding']
        for col, header in enumerate(headers, 1):
            cell = ws_findings.cell(1, col, header)
            cell.fill = header_fill
            cell.font = header_font
        
        row = 2
        for host in sorted(self.data.get('detailed_scans', {}).keys()):
            host_data = self.data['detailed_scans'][host]
            services = host_data.get('services', [])
            
            for service in services:
                scripts = service.get('scripts', [])
                for script in scripts:
                    ws_findings.cell(row, 1, host)
                    ws_findings.cell(row, 2, service.get('port', 'N/A'))
                    ws_findings.cell(row, 3, script.get('id', 'N/A'))
                    ws_findings.cell(row, 4, script.get('output', 'N/A'))
                    row += 1
        
        ws_findings.column_dimensions['A'].width = 20
        ws_findings.column_dimensions['B'].width = 10
        ws_findings.column_dimensions['C'].width = 25
        ws_findings.column_dimensions['D'].width = 60
        
        # Save workbook
        output_file = self.output_dir / f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        wb.save(output_file)
        
        print(f"[+] XLSX report saved: {output_file}")
        return output_file


def main():
    parser = argparse.ArgumentParser(
        description='Generate DOCX and XLSX reports from scan JSON results'
    )
    parser.add_argument('json_file', help='JSON scan results file')
    parser.add_argument('--docx', action='store_true', help='Generate DOCX report')
    parser.add_argument('--xlsx', action='store_true', help='Generate XLSX report')
    parser.add_argument('--all', action='store_true', help='Generate all report types')
    
    args = parser.parse_args()
    
    if not Path(args.json_file).exists():
        print(f"[!] Error: File not found: {args.json_file}")
        sys.exit(1)
    
    generator = EnhancedReportGenerator(args.json_file)
    
    generated = []
    
    if args.all or args.docx:
        docx_file = generator.generate_docx_report()
        if docx_file:
            generated.append(docx_file)
    
    if args.all or args.xlsx:
        xlsx_file = generator.generate_xlsx_report()
        if xlsx_file:
            generated.append(xlsx_file)
    
    if not (args.docx or args.xlsx or args.all):
        print("[!] No report type specified. Use --docx, --xlsx, or --all")
        parser.print_help()
        sys.exit(1)
    
    if generated:
        print("\n[+] Report generation complete!")
        print(f"[+] Generated {len(generated)} report(s)")
    else:
        print("\n[!] No reports were generated. Check dependencies.")


if __name__ == "__main__":
    main()
